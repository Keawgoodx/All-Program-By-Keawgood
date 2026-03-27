"""
╔══════════════════════════════════════════════════════════════════════════╗
║              Keawgood Universe — All-in-One Hub                         ║
║  รวมทุกโปรแกรมไว้ในไฟล์เดียว ทุกโปรแกรมใช้งานได้ 100%               ║
╚══════════════════════════════════════════════════════════════════════════╝
"""

# ══════════════════════════════════════════════════════════
#  SHARED IMPORTS
# ══════════════════════════════════════════════════════════
import os
import re
import shutil
import threading
import concurrent.futures
import json
import time
import random
from pathlib import Path
from datetime import datetime
from urllib.parse import urljoin

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import tkinter.font as tkfont

import customtkinter as ctk

# ══════════════════════════════════════════════════════════════════════════════
#
#  ██████╗ ██╗   ██╗    ██╗  ██╗███████╗ █████╗ ██╗    ██╗ ██████╗  ██████╗  ██████╗ ██████╗
#  ██╔══██╗╚██╗ ██╔╝    ██║ ██╔╝██╔════╝██╔══██╗██║    ██║██╔════╝ ██╔═══██╗██╔═══██╗██╔══██╗
#  ██████╔╝ ╚████╔╝     █████╔╝ █████╗  ███████║██║ █╗ ██║██║  ███╗██║   ██║██║   ██║██║  ██║
#  ██╔══██╗  ╚██╔╝      ██╔═██╗ ██╔══╝  ██╔══██║██║███╗██║██║   ██║██║   ██║██║   ██║██║  ██║
#  ██████╔╝   ██║       ██║  ██╗███████╗██║  ██║╚███╔███╔╝╚██████╔╝╚██████╔╝╚██████╔╝██████╔╝
#  ╚═════╝    ╚═╝       ╚═╝  ╚═╝╚══════╝╚═╝  ╚═╝ ╚══╝╚══╝  ╚═════╝  ╚═════╝  ╚═════╝╚═════╝
#
#  by_keawgood.py — Thai Novel File Manager
# ══════════════════════════════════════════════════════════════════════════════

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import fitz  # PyMuPDF
    PDF_OK = True
except ImportError:
    PDF_OK = False

# ── Colors ──
BK_ACCENT   = ("#6366F1", "#7C6BFF")
BK_ACCENT2  = ("#4F46E5", "#A78BFA")
BK_BG_DARK  = ("#F1F5F9", "#0F0F1A")
BK_BG_CARD  = ("#FFFFFF", "#1A1A2E")
BK_BG_INPUT = ("#E2E8F0", "#252540")
BK_FG_TEXT  = ("#1E293B", "#E2E8F0")
BK_FG_MUTED = ("#64748B", "#94A3B8")
BK_SUCCESS  = ("#10B981", "#34D399")
BK_ERROR    = ("#EF4444", "#F87171")
BK_WARN     = ("#F59E0B", "#FBBF24")

CHAPTER_PATTERN = re.compile(r'^(?:ตอนที่|ตอน|บทที่|บท|Chapter)\s*\d+', re.IGNORECASE)


def bk_natural_key(name: str):
    nums = re.findall(r'\d+', name)
    return int(nums[0]) if nums else 0


def bk_read_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        if not PDF_OK:
            raise ImportError("PyMuPDF ไม่ได้ติดตั้ง — pip install PyMuPDF")
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    if ext == ".docx":
        if not DOCX_OK:
            raise ImportError("python-docx ไม่ได้ติดตั้ง — pip install python-docx")
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs)
    for enc in ("utf-8-sig", "utf-8", "tis-620", "cp874"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"อ่านไฟล์ไม่ได้: {path}")


def bk_write_text(path: str, text: str):
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8-sig") as f:
        f.write(text)


def bk_write_docx(path: str, text: str):
    if not DOCX_OK:
        raise ImportError("python-docx ไม่ได้ติดตั้ง")
    doc = DocxDocument()
    for line in text.splitlines():
        line = line.strip('\u200b\ufeff \t')
        if CHAPTER_PATTERN.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc.save(path)


def bk_clean_text(text: str) -> str:
    text = text.replace('\u200b', '').replace('\ufeff', '')
    lines = [l.rstrip() for l in text.splitlines()]
    return "\n".join(lines)


def bk_safe_filename(name: str, max_len: int = 100) -> str:
    name = re.sub(r'[\\/*?:"<>|]', "", name).strip()
    return name[:max_len].strip() or "unnamed"


def bk_list_files(folder: str, exts=(".txt", ".md", ".docx", ".pdf")):
    files = [
        f for f in os.listdir(folder)
        if os.path.isfile(os.path.join(folder, f))
        and os.path.splitext(f)[1].lower() in exts
    ]
    files.sort(key=bk_natural_key)
    return files


def bk_merge_files(src_folder, out_file, log):
    files = bk_list_files(src_folder)
    if not files:
        log("❌ ไม่พบไฟล์ในโฟลเดอร์"); return
    log(f"📂 พบไฟล์ {len(files)} ไฟล์ กำลังรวม...")
    chunks = []
    for fn in files:
        path = os.path.join(src_folder, fn)
        if os.path.abspath(path) == os.path.abspath(out_file):
            continue
        try:
            chunks.append(bk_clean_text(bk_read_text(path)))
            log(f"   ✅ {fn}")
        except Exception as e:
            log(f"   ⚠️ {fn} — {e}")
    combined = "\n\n---\n\n".join(chunks)
    if os.path.splitext(out_file)[1].lower() == ".docx":
        bk_write_docx(out_file, combined)
    else:
        bk_write_text(out_file, combined)
    log(f"\n🎉 รวมไฟล์เสร็จแล้ว → {out_file}")


def bk_split_to_chapters(src_path_or_folder, out_folder, out_ext, is_folder, prefix, log):
    paths = ([os.path.join(src_path_or_folder, f) for f in bk_list_files(src_path_or_folder)]
             if is_folder else [src_path_or_folder])
    os.makedirs(out_folder, exist_ok=True)
    
    global_chapter_count = 1
    total_files = 0
    
    for path in paths:
        log(f"📖 อ่านไฟล์เพื่อแยกตอน: {os.path.basename(path)}")
        try:
            raw = bk_clean_text(bk_read_text(path))
        except Exception as e:
            log(f"   ❌ {e}"); continue
            
        lines = raw.splitlines()
        chapter_title = "บทนำ_และเนื้อหาเริ่มต้น"
        chapter_lines = []
        
        def save_ch(title, content_lines):
            nonlocal global_chapter_count, total_files
            if not content_lines: return
            text = "\n".join(content_lines).strip()
            if not text: return
            
            p = prefix.strip()
            if p:
                if "[n]" in p:
                    fname = p.replace("[n]", f"{global_chapter_count:03d}")
                else:
                    fname = f"{p} {global_chapter_count:03d}"
            else:
                fname = bk_safe_filename(title)
                
            fname += (out_ext or ".txt")
            fpath = os.path.join(out_folder, fname)
            
            if out_ext == ".docx":
                bk_write_docx(fpath, text)
            else:
                bk_write_text(fpath, text)
                
            total_files += 1
            global_chapter_count += 1
            log(f"   💾 บันทึก: {fname}")

        for line in lines:
            cl = line.strip('\u200b\ufeff \t\r')
            if not cl:
                chapter_lines.append(""); continue
            if CHAPTER_PATTERN.match(cl):
                save_ch(chapter_title, chapter_lines)
                chapter_title = cl
                chapter_lines = [cl]
                log(f"   🔖 เจอตอน: {cl[:60]}")
            else:
                chapter_lines.append(cl)
        save_ch(chapter_title, chapter_lines)
        
    log(f"\n🎉 แยกตอนเสร็จสิ้น! ได้ทั้งหมด {total_files} ไฟล์ → {out_folder}")


def bk_batch_merge_files(src_folder, out_folder, batch_size_str, prefix, out_ext, log):
    try:
        batch_size = int(batch_size_str)
        if batch_size <= 0: raise ValueError
    except ValueError:
        log("❌ กรุณาระบุจำนวนไฟล์ต่อ 1 กลุ่มเป็นตัวเลขที่มากกว่า 0"); return

    raw_files = bk_list_files(src_folder)
    files = []
    
    # ตัวกรองอัจฉริยะ: ป้องกันการดึงไฟล์ที่รวมทุกตอนหรือไฟล์ที่จับกลุ่มไปแล้วมาอ่านซ้ำ
    for f in raw_files:
        # 1. ข้ามไฟล์ที่ชื่อมี " - " (พวกไฟล์กลุ่มที่อาจหลงอยู่ในโฟลเดอร์)
        if " - " in f:
            continue
        # 2. ข้ามไฟล์ที่ไม่มีตัวเลขในชื่อเลย (พวก 'รวมไฟล์ทั้งหมด.txt', 'บทนำ.txt' ฯลฯ)
        if not re.search(r'\d+', f):
            continue
        files.append(f)

    if not files:
        log("❌ ไม่พบไฟล์ตอนย่อยในโฟลเดอร์ต้นทาง (เช็คว่ามีตัวเลขในชื่อและไม่มีเครื่องหมาย '-')"); return
        
    os.makedirs(out_folder, exist_ok=True)
    total_groups = 0
    
    for i in range(0, len(files), batch_size):
        batch = files[i:i+batch_size]
        
        # ดึงเลขตอนจากชื่อไฟล์แรกและไฟล์สุดท้ายของกลุ่มนี้จริงๆ เพื่อตั้งชื่อให้เป๊ะ
        start_num = bk_natural_key(batch[0])
        end_num = bk_natural_key(batch[-1])
        
        chunks = []
        for fn in batch:
            path = os.path.join(src_folder, fn)
            log(f"📖 กำลังรวมไฟล์: {fn}")
            try:
                chunks.append(bk_clean_text(bk_read_text(path)))
            except Exception as e:
                log(f"   ❌ ข้ามไฟล์ {fn} เนื่องจาก: {e}")
                
        if not chunks: continue
        
        # เชื่อมไฟล์ในกลุ่มด้วยตัวแบ่ง
        combined = "\n\n---\n\n".join(chunks)
        
        p = prefix.strip()
        if p:
            fname = f"{p} {start_num:03d} - {end_num:03d}{out_ext}"
        else:
            fname = f"{start_num:03d}-{end_num:03d}{out_ext}"
            
        fpath = os.path.join(out_folder, fname)
        
        if out_ext == ".docx":
            bk_write_docx(fpath, combined)
        else:
            bk_write_text(fpath, combined)
            
        total_groups += 1
        log(f"   📦 บันทึกกลุ่มไฟล์สำเร็จ: {fname}")
        
    log(f"\n🎉 รวมกลุ่มเสร็จสิ้น! ได้ผลลัพธ์ทั้งหมด {total_groups} ไฟล์ → {out_folder}")


BK_CONVERSIONS = [
    (".txt", ".md",   "TXT → MD"),
    (".md",  ".txt",  "MD → TXT"),
    (".docx",".txt",  "DOCX → TXT"),
    (".txt", ".docx", "TXT → DOCX"),
]


def bk_convert_files(src_folder, out_folder, from_ext, to_ext, log):
    files = [f for f in os.listdir(src_folder)
             if f.lower().endswith(from_ext)
             and os.path.isfile(os.path.join(src_folder, f))]
    if not files:
        log(f"❌ ไม่พบไฟล์ {from_ext} ในโฟลเดอร์"); return
    os.makedirs(out_folder, exist_ok=True)
    count = 0
    for fn in files:
        src = os.path.join(src_folder, fn)
        base = os.path.splitext(fn)[0]
        try:
            if from_ext == ".docx" and to_ext == ".txt":
                bk_write_text(os.path.join(out_folder, base + to_ext), bk_clean_text(bk_read_text(src)))
            elif from_ext == ".txt" and to_ext == ".docx":
                bk_write_docx(os.path.join(out_folder, base + to_ext), bk_clean_text(bk_read_text(src)))
            else:
                shutil.copy2(src, os.path.join(out_folder, base + to_ext))
            count += 1
            log(f"   ✅ {fn} → {base + to_ext}")
        except Exception as e:
            log(f"   ❌ {fn}: {e}")
    log(f"\n🎉 แปลงไฟล์เสร็จ {count} ไฟล์ → {out_folder}")


# ── UI Widgets for ByKeawgood ──

class BK_LogBox(ctk.CTkTextbox):
    def __init__(self, master, **kw):
        super().__init__(master, font=("Consolas", 12), text_color=BK_FG_TEXT,
                         fg_color=BK_BG_INPUT, corner_radius=10, wrap="word", **kw)
        self.configure(state="disabled")

    def write(self, msg: str):
        self.configure(state="normal")
        self.insert("end", msg + "\n")
        self.see("end")
        self.configure(state="disabled")
        self.update_idletasks()

    def clear(self):
        self.configure(state="normal")
        self.delete("1.0", "end")
        self.configure(state="disabled")


def bk_row(parent, label, var, browse_cmd, btn_label="📂 เลือก"):
    frame = ctk.CTkFrame(parent, fg_color="transparent")
    frame.pack(fill="x", pady=4)
    ctk.CTkLabel(frame, text=label, font=("Kanit", 13), text_color=BK_FG_MUTED,
                 width=160, anchor="w").pack(side="left")
    ctk.CTkEntry(frame, textvariable=var, font=("Kanit", 12),
                 fg_color=BK_BG_INPUT, border_color=BK_ACCENT,
                 corner_radius=8).pack(side="left", fill="x", expand=True, padx=(6, 6))
    ctk.CTkButton(frame, text=btn_label, command=browse_cmd,
                  fg_color=BK_ACCENT, hover_color=BK_ACCENT2,
                  font=("Kanit", 12), width=100, corner_radius=8).pack(side="left")


def bk_section_title(parent, text):
    ctk.CTkLabel(parent, text=text, font=("Kanit", 17, "bold"),
                 text_color=BK_ACCENT2).pack(anchor="w", pady=(14, 2))
    ctk.CTkFrame(parent, height=2, fg_color=BK_ACCENT, corner_radius=1).pack(fill="x", pady=(0, 10))


def bk_run_btn(parent, text, cmd):
    return ctk.CTkButton(parent, text=text, command=cmd,
                         fg_color=BK_ACCENT, hover_color=BK_ACCENT2,
                         font=("Kanit", 14, "bold"), height=44, corner_radius=10)


class BK_MergeTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var = ctk.StringVar()
        self.out_var = ctk.StringVar()
        self.out_name_var = ctk.StringVar(value="รวมไฟล์ทั้งหมด.txt")
        self._build()

    def _build(self):
        bk_section_title(self, "  🗂️  รวมไฟล์ทั้งหมด (Merge All)")
        bk_row(self, "โฟลเดอร์ต้นทาง:", self.src_var,
               lambda: self.src_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์ต้นทาง")))
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์ปลายทาง")))
        row = ctk.CTkFrame(self, fg_color="transparent")
        row.pack(fill="x", pady=4)
        ctk.CTkLabel(row, text="ชื่อไฟล์ผลลัพธ์:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        ctk.CTkEntry(row, textvariable=self.out_name_var, font=("Kanit", 12),
                     fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8
                     ).pack(side="left", fill="x", expand=True, padx=6)
        bk_run_btn(self, "▶  เริ่มรวมไฟล์", self._run).pack(pady=18)

    def _run(self):
        src = self.src_var.get().strip()
        out_dir = self.out_var.get().strip()
        name = self.out_name_var.get().strip()
        if not src or not out_dir or not name:
            messagebox.showwarning("By Keawgood", "กรุณากรอกข้อมูลให้ครบก่อนนะครับ"); return
        self.log.clear()
        threading.Thread(target=bk_merge_files,
                         args=(src, os.path.join(out_dir, name), self.log.write),
                         daemon=True).start()


class BK_SplitTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var  = ctk.StringVar()
        self.out_var  = ctk.StringVar()
        self.mode_var = ctk.StringVar(value="file")
        self.prefix_var = ctk.StringVar(value="ตอนที่ [n]")
        self.ext_var  = ctk.StringVar(value=".txt")
        self._build()

    def _build(self):
        bk_section_title(self, "  ✂️  แยกไฟล์ (ทีละ 1 ตอน)")
        mode_f = ctk.CTkFrame(self, fg_color="transparent")
        mode_f.pack(fill="x", pady=6)
        ctk.CTkLabel(mode_f, text="โหมด:", font=("Kanit", 13), text_color=BK_FG_MUTED,
                     width=160, anchor="w").pack(side="left")
        ctk.CTkRadioButton(mode_f, text="ไฟล์เดียว (ไฟล์รวมหลายตอน)", variable=self.mode_var, value="file",
                           font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=10)
        ctk.CTkRadioButton(mode_f, text="ทั้งโฟลเดอร์", variable=self.mode_var, value="folder",
                           font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=10)
        bk_row(self, "ต้นทาง:", self.src_var, self._browse_src)
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์สำหรับเซฟไฟล์ที่แยกแล้ว")))
        
        opt_f = ctk.CTkFrame(self, fg_color="transparent")
        opt_f.pack(fill="x", pady=4)
        ctk.CTkLabel(opt_f, text="ตั้งชื่อไฟล์ (ใช้ [n] แทนเลขตอน):", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        ctk.CTkEntry(opt_f, textvariable=self.prefix_var, font=("Kanit", 12),
                     fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8).pack(side="left", fill="x", expand=True, padx=(6, 6))

        ext_f = ctk.CTkFrame(self, fg_color="transparent")
        ext_f.pack(fill="x", pady=4)
        ctk.CTkLabel(ext_f, text="บันทึกเป็นนามสกุล:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        for e in (".txt", ".md", ".docx"):
            ctk.CTkRadioButton(ext_f, text=e, variable=self.ext_var, value=e,
                               font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=8)
        bk_run_btn(self, "▶  เริ่มแยกตอน", self._run).pack(pady=18)

    def _browse_src(self):
        if self.mode_var.get() == "file":
            p = filedialog.askopenfilename(filetypes=[("ไฟล์นิยาย", "*.txt *.md *.docx *.pdf")])
        else:
            p = filedialog.askdirectory()
        if p: self.src_var.set(p)

    def _run(self):
        src = self.src_var.get().strip()
        out = self.out_var.get().strip()
        if not src or not out:
            messagebox.showwarning("By Keawgood", "กรุณากรอกข้อมูลให้ครบก่อนนะครับ"); return
        self.log.clear()
        threading.Thread(target=bk_split_to_chapters,
                         args=(src, out, self.ext_var.get(), self.mode_var.get() == "folder", self.prefix_var.get(), self.log.write),
                         daemon=True).start()


class BK_BatchMergeTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var  = ctk.StringVar()
        self.out_var  = ctk.StringVar()
        self.batch_var = ctk.StringVar(value="5")
        self.prefix_var = ctk.StringVar(value="Chapter")
        self.ext_var  = ctk.StringVar(value=".txt")
        self._build()

    def _build(self):
        bk_section_title(self, "  📚  รวมกลุ่ม (ทีละ N ไฟล์)")
        
        bk_row(self, "โฟลเดอร์ต้นทาง:", self.src_var,
               lambda: self.src_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์ที่มีไฟล์ย่อย")))
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์บันทึก")))
        
        opt_f = ctk.CTkFrame(self, fg_color="transparent")
        opt_f.pack(fill="x", pady=4)
        
        ctk.CTkLabel(opt_f, text="จำนวนไฟล์ต่อ 1 กลุ่ม:", font=("Kanit", 13), text_color=BK_FG_MUTED, width=160, anchor="w").grid(row=0, column=0, pady=4, sticky="w")
        ctk.CTkEntry(opt_f, textvariable=self.batch_var, font=("Kanit", 12), width=80, fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8).grid(row=0, column=1, pady=4, sticky="w")
        
        ctk.CTkLabel(opt_f, text="ตั้งชื่อไฟล์ (เช่น Chapter):", font=("Kanit", 13), text_color=BK_FG_MUTED, width=160, anchor="w").grid(row=1, column=0, pady=4, sticky="w")
        ctk.CTkEntry(opt_f, textvariable=self.prefix_var, font=("Kanit", 12), width=180, fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8).grid(row=1, column=1, pady=4, sticky="w")

        ext_f = ctk.CTkFrame(self, fg_color="transparent")
        ext_f.pack(fill="x", pady=4)
        ctk.CTkLabel(ext_f, text="บันทึกเป็นนามสกุล:", font=("Kanit", 13), text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        for e in (".txt", ".md", ".docx"):
            ctk.CTkRadioButton(ext_f, text=e, variable=self.ext_var, value=e, font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=8)

        bk_run_btn(self, "▶  เริ่มรวมไฟล์เป็นกลุ่ม", self._run).pack(pady=18)

    def _run(self):
        src = self.src_var.get().strip()
        out = self.out_var.get().strip()
        if not src or not out:
            messagebox.showwarning("By Keawgood", "กรุณาเลือกโฟลเดอร์ให้ครบก่อนนะครับ")
            return
        self.log.clear()
        threading.Thread(target=bk_batch_merge_files,
                         args=(src, out, self.batch_var.get(), self.prefix_var.get(), self.ext_var.get(), self.log.write),
                         daemon=True).start()


class BK_ConvertTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var  = ctk.StringVar()
        self.out_var  = ctk.StringVar()
        self.mode_var = ctk.StringVar(value="0")
        self._build()

    def _build(self):
        bk_section_title(self, "  🔄  แปลงรูปแบบไฟล์ (Convert)")
        bk_row(self, "โฟลเดอร์ต้นทาง:", self.src_var,
               lambda: self.src_var.set(filedialog.askdirectory()))
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory()))
        ctk.CTkLabel(self, text="รูปแบบการแปลง:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED).pack(anchor="w", pady=(10, 4))
        grid = ctk.CTkFrame(self, fg_color="transparent")
        grid.pack(fill="x", padx=4)
        for i, (_, _, label) in enumerate(BK_CONVERSIONS):
            ctk.CTkRadioButton(grid, text=label, variable=self.mode_var, value=str(i),
                               font=("Kanit", 13, "bold"), fg_color=BK_ACCENT,
                               text_color=BK_FG_TEXT).grid(row=i//2, column=i%2, padx=20, pady=6, sticky="w")
        bk_run_btn(self, "▶  เริ่มแปลงไฟล์", self._run).pack(pady=18)

    def _run(self):
        src = self.src_var.get().strip()
        out = self.out_var.get().strip()
        idx = int(self.mode_var.get())
        from_ext, to_ext, _ = BK_CONVERSIONS[idx]
        if not src or not out:
            messagebox.showwarning("By Keawgood", "กรุณากรอกข้อมูลให้ครบก่อนนะครับ"); return
        self.log.clear()
        threading.Thread(target=bk_convert_files,
                         args=(src, out, from_ext, to_ext, self.log.write),
                         daemon=True).start()


class ByKeawgoodWindow(ctk.CTkToplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title("By Keawgood — Thai Novel Manager")
        self.geometry("940x760")
        self.minsize(820, 640)
        self.configure(fg_color=BK_BG_DARK)
        self._build()
        self.lift()
        self.focus_force()

    def _build(self):
        hdr = ctk.CTkFrame(self, fg_color=BK_BG_CARD, corner_radius=0, height=72)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text="✦  By Keawgood", font=("Kanit", 26, "bold"),
                     text_color=BK_ACCENT2).pack(side="left", padx=28)
        ctk.CTkLabel(hdr, text="Thai Novel File Manager",
                     font=("Kanit", 13), text_color=BK_FG_MUTED).pack(side="left")

        self._mode = ctk.StringVar(value="dark")
        self.mode_switch = ctk.CTkSwitch(hdr, text="🌙 Dark", variable=self._mode,
                                         onvalue="dark", offvalue="light",
                                         command=self._toggle_mode,
                                         font=("Kanit", 12), fg_color=BK_ACCENT,
                                         progress_color=BK_ACCENT2)
        self.mode_switch.pack(side="right", padx=24)

        body = ctk.CTkFrame(self, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=18, pady=14)
        body.columnconfigure(0, weight=1)
        body.rowconfigure(0, weight=1)

        tabview = ctk.CTkTabview(body, fg_color=BK_BG_CARD,
                                 segmented_button_fg_color=BK_BG_INPUT,
                                 segmented_button_selected_color=BK_ACCENT,
                                 segmented_button_selected_hover_color=BK_ACCENT2,
                                 segmented_button_unselected_color=BK_BG_INPUT,
                                 text_color=BK_FG_TEXT, corner_radius=14)
        tabview.grid(row=0, column=0, sticky="nsew", pady=(0, 10))
        
        for name in ("🗂️ รวมไฟล์ทั้งหมด", "📚 รวมกลุ่ม (ทีละ N ไฟล์)", "✂️ แยกไฟล์ (ทีละ 1 ตอน)", "🔄 แปลงไฟล์"):
            tabview.add(name)
            tabview.tab(name).configure(fg_color=BK_BG_CARD)

        log_frame = ctk.CTkFrame(body, fg_color=BK_BG_CARD, corner_radius=12)
        log_frame.grid(row=1, column=0, sticky="nsew")
        body.rowconfigure(1, weight=0, minsize=200)
        ctk.CTkLabel(log_frame, text="📋 Console Log",
                     font=("Kanit", 13, "bold"), text_color=BK_ACCENT2).pack(anchor="w", padx=14, pady=(8, 2))
        self.log = BK_LogBox(log_frame, height=160)
        self.log.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        ctk.CTkButton(log_frame, text="🗑 ล้าง Log", width=110,
                      fg_color=("#CBD5E1", "#3F3F5A"), hover_color=BK_ACCENT,
                      font=("Kanit", 11), corner_radius=8,
                      command=self.log.clear).pack(anchor="e", padx=10, pady=(0, 8))

        for cls, tab_name in [(BK_MergeTab, "🗂️ รวมไฟล์ทั้งหมด"),
                              (BK_BatchMergeTab, "📚 รวมกลุ่ม (ทีละ N ไฟล์)"),
                              (BK_SplitTab, "✂️ แยกไฟล์ (ทีละ 1 ตอน)"), 
                              (BK_ConvertTab, "🔄 แปลงไฟล์")]:
            widget = cls(tabview.tab(tab_name), self.log)
            widget.pack(fill="both", expand=True, padx=16, pady=12)

        self.log.write("👋 ยินดีต้อนรับสู่ By Keawgood!")
        self.log.write("   เลือก Tab ที่ต้องการแล้วกด ▶ เริ่ม\n")

    def _toggle_mode(self):
        mode = self._mode.get()
        ctk.set_appearance_mode(mode)
        self.mode_switch.configure(text="☀️ Light" if mode == "light" else "🌙 Dark")


# ══════════════════════════════════════════════════════════════════════════════
#
#  ██╗   ██╗ ██████╗  ██████╗ █████╗ ██████╗      ██████╗ ██████╗ ████████╗
#  ██║   ██║██╔═══██╗██╔════╝██╔══██╗██╔══██╗    ██╔═══██╗██╔══██╗╚══██╔══╝
#  ██║   ██║██║   ██║██║     ███████║██████╔╝    ██║   ██║██████╔╝   ██║
#  ╚██╗ ██╔╝██║   ██║██║     ██╔══██║██╔══██╗    ██║   ██║██╔═══╝    ██║
#   ╚████╔╝ ╚██████╔╝╚██████╗██║  ██║██████╔╝    ╚██████╔╝██║        ██║
#    ╚═══╝   ╚═════╝  ╚═════╝╚═╝  ╚═╝╚═════╝      ╚═════╝ ╚═╝        ╚═╝
#
#  vocab_optimizer.py — Novel Vocab Optimizer
# ══════════════════════════════════════════════════════════════════════════════

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    DND_OK = True
except ImportError:
    DND_OK = False


class VocabOptimizerWindow(ctk.CTkToplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title("Vocab Optimizer Pro By Keawgood - เครื่องมือจัดการคำศัพท์นิยาย")
        self.geometry("1200x800")
        self.font_family = "TH Sarabun PSK"
        self.font_size = 20
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self._build()
        self.lift()
        self.focus_force()

    def _build(self):
        # Sidebar
        self.sidebar_frame = ctk.CTkFrame(self, width=250, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(8, weight=1)

        ctk.CTkLabel(self.sidebar_frame, text="Vocab Optimizer\nPro By Keawgood",
                     font=ctk.CTkFont(size=24, weight="bold")).grid(row=0, column=0, padx=20, pady=(30, 20))

        ctk.CTkButton(self.sidebar_frame, text="📂 เปิดไฟล์ / ลากวาง", height=40,
                      command=self.open_file).grid(row=1, column=0, padx=20, pady=10)
        ctk.CTkButton(self.sidebar_frame, text="💾 บันทึกไฟล์", height=40,
                      command=self.save_file).grid(row=2, column=0, padx=20, pady=10)
        ctk.CTkButton(self.sidebar_frame, text="🗑️ เคลียร์หน้าจอ", height=40,
                      fg_color="#dc3545", hover_color="#c82333",
                      command=self.clear_text).grid(row=3, column=0, padx=20, pady=10)

        ctk.CTkLabel(self.sidebar_frame, text="การจัดการคำซ้ำ:", anchor="w").grid(row=4, column=0, padx=20, pady=(20, 0), sticky="w")
        self.option_keep = ctk.CTkOptionMenu(self.sidebar_frame,
                                             values=["เก็บอันที่ยาวที่สุด (แนะนำ)",
                                                     "เก็บอันล่าสุดที่เจอ",
                                                     "เก็บอันแรกที่เจอ"])
        self.option_keep.grid(row=5, column=0, padx=20, pady=(5, 10))

        ctk.CTkLabel(self.sidebar_frame, text="ปรับโทนสีหน้าต่าง:", anchor="w").grid(row=6, column=0, padx=20, pady=(10, 0), sticky="w")
        ctk.CTkOptionMenu(self.sidebar_frame, values=["Dark", "Light", "System"],
                          command=lambda v: ctk.set_appearance_mode(v)
                          ).grid(row=7, column=0, padx=20, pady=(5, 10))

        ctk.CTkButton(self.sidebar_frame, text="⚡ ตัดคำซ้ำ!",
                      font=ctk.CTkFont(size=20, weight="bold"), height=50,
                      fg_color="#28a745", hover_color="#218838",
                      command=self.process_text).grid(row=9, column=0, padx=20, pady=(10, 30))

        # Main area
        self.main_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(1, weight=1)
        self.main_frame.grid_rowconfigure(1, weight=1)

        toolbar_frame = ctk.CTkFrame(self.main_frame, height=40)
        toolbar_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        ctk.CTkLabel(toolbar_frame, text="ขนาดฟ้อนต์:", font=("Tahoma", 14)).pack(side="left", padx=(15, 5), pady=5)
        ctk.CTkButton(toolbar_frame, text="A-", width=40, fg_color="#6c757d", hover_color="#5a6268",
                      command=self.decrease_font).pack(side="left", padx=5, pady=5)
        ctk.CTkButton(toolbar_frame, text="A+", width=40, fg_color="#6c757d", hover_color="#5a6268",
                      command=self.increase_font).pack(side="left", padx=5, pady=5)

        self.txt_input = ctk.CTkTextbox(self.main_frame, font=(self.font_family, self.font_size), wrap="none")
        self.txt_input.grid(row=1, column=0, padx=(0, 10), sticky="nsew")
        self.txt_input.insert("1.0", "--- ลากไฟล์ .txt มาวางที่นี่ หรือกดปุ่มเปิดไฟล์ที่เมนูด้านซ้าย ---")

        self.txt_output = ctk.CTkTextbox(self.main_frame, font=(self.font_family, self.font_size), wrap="none")
        self.txt_output.grid(row=1, column=1, padx=(10, 0), sticky="nsew")
        self.txt_output.insert("1.0", "--- ผลลัพธ์จะแสดงที่นี่ ---")

        # Drag & Drop (ถ้ามี tkinterdnd2)
        if DND_OK:
            try:
                self.txt_input.drop_target_register(DND_FILES)
                self.txt_input.dnd_bind('<<Drop>>', self.handle_drop)
            except Exception:
                pass

    def increase_font(self):
        self.font_size += 2
        self._update_font()

    def decrease_font(self):
        if self.font_size > 12:
            self.font_size -= 2
            self._update_font()

    def _update_font(self):
        self.txt_input.configure(font=(self.font_family, self.font_size))
        self.txt_output.configure(font=(self.font_family, self.font_size))

    def clear_text(self):
        self.txt_input.delete("1.0", ctk.END)
        self.txt_output.delete("1.0", ctk.END)

    def handle_drop(self, event):
        filepath = event.data.replace('{', '').replace('}', '').strip()
        if filepath.lower().endswith('.txt'):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    self.txt_input.delete("1.0", ctk.END)
                    self.txt_input.insert(ctk.END, f.read())
            except Exception as e:
                messagebox.showerror("Error", f"ไม่สามารถอ่านไฟล์ได้:\n{str(e)}")
        else:
            messagebox.showwarning("Warning", "กรุณาลากไฟล์นามสกุล .txt เท่านั้นครับ")

    def open_file(self):
        filepath = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if filepath:
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    self.txt_input.delete("1.0", ctk.END)
                    self.txt_input.insert(ctk.END, f.read())
            except Exception as e:
                messagebox.showerror("Error", f"ไม่สามารถอ่านไฟล์ได้:\n{str(e)}")

    def save_file(self):
        filepath = filedialog.asksaveasfilename(defaultextension=".txt",
                                                filetypes=[("Text Files", "*.txt")])
        if filepath:
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(self.txt_output.get("1.0", ctk.END))
                messagebox.showinfo("Success", "บันทึกไฟล์เรียบร้อยแล้ว!")
            except Exception as e:
                messagebox.showerror("Error", f"ไม่สามารถบันทึกไฟล์ได้:\n{str(e)}")

    def process_text(self):
        text = self.txt_input.get("1.0", ctk.END)
        lines = text.split('\n')
        vocab_dict = {}
        keep_mode = self.option_keep.get()
        for i, line in enumerate(lines):
            original_line = line.strip()
            if not original_line or original_line.startswith('---'):
                continue
            if original_line.startswith('[') and original_line.endswith(']'):
                vocab_dict[f"__HEADER_{i}"] = {'line': original_line, 'order': i, 'is_header': True}
                continue
            clean_line = re.sub(r'^[\-\s]+', '', original_line)
            parts = re.split(r'[\s=/]+', clean_line, maxsplit=1)
            if parts:
                key = parts[0].strip()
                if not key: continue
                if key in vocab_dict and not vocab_dict[key].get('is_header', False):
                    if keep_mode == "เก็บอันที่ยาวที่สุด (แนะนำ)":
                        if len(original_line) > len(vocab_dict[key]['line']):
                            vocab_dict[key]['line'] = original_line
                    elif keep_mode == "เก็บอันล่าสุดที่เจอ":
                        vocab_dict[key]['line'] = original_line
                else:
                    vocab_dict[key] = {'line': original_line, 'order': i, 'is_header': False}
        sorted_vocabs = sorted(vocab_dict.values(), key=lambda x: x['order'])
        result_lines = []
        for v in sorted_vocabs:
            if v.get('is_header', False) and result_lines and result_lines[-1] != "":
                result_lines.append("")
            result_lines.append(v['line'])
        self.txt_output.delete("1.0", ctk.END)
        self.txt_output.insert(ctk.END, '\n'.join(result_lines))


# ══════════════════════════════════════════════════════════════════════════════
#
#   █████╗ ██╗   ██╗██████╗ ██╗ ██████╗     ██████╗ ██╗   ██╗
#  ██╔══██╗██║   ██║██╔══██╗██║██╔═══██╗    ██╔══██╗╚██╗ ██╔╝
#  ███████║██║   ██║██║  ██║██║██║   ██║    ██████╔╝ ╚████╔╝
#  ██╔══██║██║   ██║██║  ██║██║██║   ██║    ██╔══██╗  ╚██╔╝
#  ██║  ██║╚██████╔╝██████╔╝██║╚██████╔╝    ██████╔╝   ██║
#  ╚═╝  ╚═╝ ╚═════╝ ╚═════╝ ╚═╝ ╚═════╝    ╚═════╝    ╚═╝
#
#  Audio_By_Keawgood.py — Batch Audio to Video
# ══════════════════════════════════════════════════════════════════════════════

AUDIO_LANGUAGES = {
    "TH": {
        "app_name": "Audio By Keawgood",
        "menu": "เมนูหลัก",
        "appearance_settings": "ตั้งค่าลักษณะ",
        "font_size": "ขนาดตัวอักษร:",
        "theme": "โหมดสี:",
        "lang_btn": "English",
        "select_audio": "1. เลือกไฟล์เสียง (.m4a)",
        "select_image": "2. เลือกรูปภาพ (.jpg/.png)",
        "select_output": "3. เลือกโฟลเดอร์ปลายทาง",
        "merge_settings": "4. ตั้งค่าการรวมและชื่อไฟล์",
        "filename_label": "ตั้งชื่อไฟล์ผลลัพธ์ (ไม่ต้องใส่ .mp4):",
        "chunk_label": "ระบุจำนวนคลิปที่จะรวมเป็น 1 ไฟล์:",
        "start": "เริ่มสร้างวิดีโอ",
        "status_idle": "สถานะ: พร้อมทำงาน",
        "status_done": "เสร็จสมบูรณ์!",
        "error_select": "กรุณาเลือกไฟล์และโฟลเดอร์ให้ครบถ้วน",
        "error_number": "กรุณาพิมพ์ตัวเลขจำนวนเต็มที่มากกว่า 0",
        "processing": "กำลังเรนเดอร์ไฟล์ที่ {current}/{total}...",
        "files_selected": "เลือกแล้ว {count} ไฟล์",
        "img_selected": "เลือกรูปภาพแล้ว",
        "output_selected": "บันทึกที่: {path}"
    },
    "EN": {
        "app_name": "Audio By Keawgood",
        "menu": "Main Menu",
        "appearance_settings": "Appearance Settings",
        "font_size": "Font Size:",
        "theme": "Theme Mode:",
        "lang_btn": "ภาษาไทย",
        "select_audio": "1. Select Audio (.m4a)",
        "select_image": "2. Select Image (.jpg/.png)",
        "select_output": "3. Select Output Folder",
        "merge_settings": "4. Merge & Naming Settings",
        "filename_label": "Output Filename (no .mp4):",
        "chunk_label": "Enter number of clips to merge into 1 file:",
        "start": "Start Rendering",
        "status_idle": "Status: Ready",
        "status_done": "Completed!",
        "error_select": "Please select all required files and folder",
        "error_number": "Please enter a valid number greater than 0",
        "processing": "Rendering file {current}/{total}...",
        "files_selected": "{count} files selected",
        "img_selected": "Image selected",
        "output_selected": "Save to: {path}"
    }
}


class AudioByKeawgoodWindow(ctk.CTkToplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title("Audio By Keawgood - Batch Audio to Video")
        self.geometry("850x650")
        self.current_lang = "TH"
        self.audio_paths = []
        self.image_path = ""
        self.output_dir = ""
        self.chunk_size = 1
        self.current_font_size = 14
        self.font = ctk.CTkFont(family="TH Sarabun PSK", size=self.current_font_size)
        self.font_bold = ctk.CTkFont(family="TH Sarabun PSK", size=self.current_font_size, weight="bold")
        self._setup_ui()
        self.lift()
        self.focus_force()

    def _L(self):
        return AUDIO_LANGUAGES[self.current_lang]

    def _setup_ui(self):
        L = self._L()

        self.sidebar_frame = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sidebar_frame.pack(side="left", fill="y")

        ctk.CTkLabel(self.sidebar_frame, text=L["app_name"],
                     font=("TH Sarabun PSK", 22, "bold")).pack(pady=(30, 10), padx=20)
        self.lbl_menu = ctk.CTkLabel(self.sidebar_frame, text=L["menu"], text_color="gray", font=self.font)
        self.lbl_menu.pack(pady=(10, 5))

        self.lang_btn = ctk.CTkButton(self.sidebar_frame, text=L["lang_btn"],
                                      command=self._toggle_language, fg_color="transparent",
                                      border_width=1, text_color=("gray10", "#DCE4EE"), font=self.font)
        self.lang_btn.pack(pady=10, padx=20)

        frame_app = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        frame_app.pack(pady=(20, 10), fill="x", padx=10)
        self.lbl_appearance = ctk.CTkLabel(frame_app, text=L["appearance_settings"], text_color="gray", font=self.font_bold)
        self.lbl_appearance.pack(pady=(10, 5))
        self.lbl_font_size = ctk.CTkLabel(frame_app, text=L["font_size"], font=self.font)
        self.lbl_font_size.pack(pady=(5, 0))
        self.lbl_font_size_display = ctk.CTkLabel(frame_app, text=f"{self.current_font_size} pt",
                                                   text_color="#3498db", font=self.font_bold)
        self.lbl_font_size_display.pack(pady=(0, 5))
        self.slider_font_size = ctk.CTkSlider(frame_app, from_=10, to=30, number_of_steps=20,
                                               command=self._update_font_size)
        self.slider_font_size.set(self.current_font_size)
        self.slider_font_size.pack(pady=(0, 15), padx=20)
        self.lbl_theme = ctk.CTkLabel(frame_app, text=L["theme"], font=self.font)
        self.lbl_theme.pack(pady=(5, 0))
        self.appearance_mode_menu = ctk.CTkOptionMenu(frame_app, values=["Dark", "Light"],
                                                       command=lambda v: ctk.set_appearance_mode(v),
                                                       font=self.font)
        self.appearance_mode_menu.pack(pady=10, padx=20)

        self.main_frame = ctk.CTkScrollableFrame(self, corner_radius=15)
        self.main_frame.pack(side="right", fill="both", expand=True, padx=20, pady=20)

        frame_audio = ctk.CTkFrame(self.main_frame, fg_color=("gray90", "gray15"))
        frame_audio.pack(fill="x", pady=10, padx=10)
        self.btn_audio = ctk.CTkButton(frame_audio, text=L["select_audio"], command=self._select_audio, font=self.font)
        self.btn_audio.pack(side="left", pady=15, padx=20)
        self.lbl_audio_status = ctk.CTkLabel(frame_audio, text="0 files", font=self.font)
        self.lbl_audio_status.pack(side="left", padx=10)

        frame_image = ctk.CTkFrame(self.main_frame, fg_color=("gray90", "gray15"))
        frame_image.pack(fill="x", pady=10, padx=10)
        self.btn_image = ctk.CTkButton(frame_image, text=L["select_image"],
                                       fg_color="#2ecc71", hover_color="#27ae60",
                                       command=self._select_image, font=self.font)
        self.btn_image.pack(side="left", pady=15, padx=20)
        self.lbl_image_status = ctk.CTkLabel(frame_image, text="", font=self.font)
        self.lbl_image_status.pack(side="left", padx=10)

        frame_output = ctk.CTkFrame(self.main_frame, fg_color=("gray90", "gray15"))
        frame_output.pack(fill="x", pady=10, padx=10)
        self.btn_output = ctk.CTkButton(frame_output, text=L["select_output"],
                                        fg_color="#e67e22", hover_color="#d35400",
                                        command=self._select_output, font=self.font)
        self.btn_output.pack(side="left", pady=15, padx=20)
        self.lbl_output_status = ctk.CTkLabel(frame_output, text="", wraplength=400,
                                              justify="left", font=self.font)
        self.lbl_output_status.pack(side="left", padx=10, pady=10)

        frame_settings = ctk.CTkFrame(self.main_frame, fg_color=("gray90", "gray15"))
        frame_settings.pack(fill="x", pady=10, padx=10)
        self.lbl_settings = ctk.CTkLabel(frame_settings, text=L["merge_settings"], font=self.font_bold)
        self.lbl_settings.pack(pady=(10, 5))
        self.lbl_filename = ctk.CTkLabel(frame_settings, text=L["filename_label"], font=self.font)
        self.lbl_filename.pack(pady=(5, 0))
        self.entry_filename = ctk.CTkEntry(frame_settings, placeholder_text="เช่น นิยายเรื่อง...",
                                           width=300, justify="center", font=self.font)
        self.entry_filename.insert(0, "Audio_Output")
        self.entry_filename.pack(pady=(0, 15))
        self.lbl_chunk = ctk.CTkLabel(frame_settings, text=L["chunk_label"], font=self.font)
        self.lbl_chunk.pack(pady=(0, 0))
        self.entry_chunk = ctk.CTkEntry(frame_settings, placeholder_text="เช่น 10, 50 หรือจำนวนทั้งหมด",
                                        width=300, justify="center", font=self.font)
        self.entry_chunk.insert(0, "1")
        self.entry_chunk.pack(pady=(0, 15))

        self.progress_bar = ctk.CTkProgressBar(self.main_frame)
        self.progress_bar.set(0)
        self.progress_bar.pack(pady=(20, 10), fill="x", padx=20)
        self.lbl_status = ctk.CTkLabel(self.main_frame, text=L["status_idle"],
                                       text_color="#3498db", font=self.font)
        self.lbl_status.pack()
        self.btn_start = ctk.CTkButton(self.main_frame, text=L["start"],
                                       height=50, fg_color="#e74c3c", hover_color="#c0392b",
                                       command=self._start_process_thread, font=self.font_bold)
        self.btn_start.pack(pady=20, padx=20, fill="x")

    def _update_font_size(self, value):
        self.current_font_size = int(value)
        self.font.configure(size=self.current_font_size)
        self.font_bold.configure(size=self.current_font_size)
        self.lbl_font_size_display.configure(text=f"{self.current_font_size} pt")

    def _toggle_language(self):
        self.current_lang = "EN" if self.current_lang == "TH" else "TH"
        L = self._L()
        self.lbl_menu.configure(text=L["menu"])
        self.lbl_appearance.configure(text=L["appearance_settings"])
        self.lbl_font_size.configure(text=L["font_size"])
        self.lbl_font_size_display.configure(text=f"{self.current_font_size} pt")
        self.lbl_theme.configure(text=L["theme"])
        self.btn_audio.configure(text=L["select_audio"])
        self.btn_image.configure(text=L["select_image"])
        self.btn_output.configure(text=L["select_output"])
        self.lbl_settings.configure(text=L["merge_settings"])
        self.lbl_filename.configure(text=L["filename_label"])
        self.lbl_chunk.configure(text=L["chunk_label"])
        self.btn_start.configure(text=L["start"])
        self.lang_btn.configure(text=L["lang_btn"])
        self.lbl_status.configure(text=L["status_idle"])
        if self.audio_paths:
            self.lbl_audio_status.configure(text=L["files_selected"].format(count=len(self.audio_paths)))
        if self.image_path:
            self.lbl_image_status.configure(text=L["img_selected"])
        if self.output_dir:
            self.lbl_output_status.configure(text=L["output_selected"].format(path=self.output_dir))

    def _select_audio(self):
        files = filedialog.askopenfilenames(filetypes=[("Audio files", "*.m4a *.mp3 *.wav")])
        if files:
            self.audio_paths = sorted(list(files))
            total_files = len(self.audio_paths)
            self.lbl_audio_status.configure(
                text=self._L()["files_selected"].format(count=total_files))
            self.entry_chunk.delete(0, 'end')
            self.entry_chunk.insert(0, str(total_files))

    def _select_image(self):
        file = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png")])
        if file:
            self.image_path = file
            self.lbl_image_status.configure(text=self._L()["img_selected"])

    def _select_output(self):
        path = filedialog.askdirectory()
        if path:
            self.output_dir = path
            self.lbl_output_status.configure(
                text=self._L()["output_selected"].format(path=path))

    def _start_process_thread(self):
        if not self.audio_paths or not self.image_path or not self.output_dir:
            messagebox.showwarning("Error", self._L()["error_select"]); return
        try:
            val = int(self.entry_chunk.get())
            if val <= 0: raise ValueError
            self.chunk_size = val
        except ValueError:
            messagebox.showwarning("Error", self._L()["error_number"]); return
        self.btn_start.configure(state="disabled")
        self.entry_chunk.configure(state="disabled")
        self.entry_filename.configure(state="disabled")
        self.slider_font_size.configure(state="disabled")
        threading.Thread(target=self._process_files, daemon=True).start()

    def _process_files(self):
        try:
            import moviepy.editor as mpe
        except ImportError:
            messagebox.showerror("Error", "กรุณาติดตั้ง moviepy ก่อน: pip install moviepy"); return

        L = self._L()
        files = self.audio_paths
        chunk_size = self.chunk_size
        base_filename = self.entry_filename.get().strip() or "Audio_Output"
        total_chunks = (len(files) + chunk_size - 1) // chunk_size

        for i in range(0, len(files), chunk_size):
            current_chunk_index = (i // chunk_size) + 1
            chunk_files = files[i:i + chunk_size]
            self.lbl_status.configure(text=L["processing"].format(current=current_chunk_index, total=total_chunks))
            self.progress_bar.set(current_chunk_index / total_chunks)
            try:
                audio_clips = [mpe.AudioFileClip(f) for f in chunk_files]
                final_audio = mpe.concatenate_audioclips(audio_clips)
                img_clip = mpe.ImageClip(self.image_path)
                img_clip = img_clip.set_duration(final_audio.duration)
                img_clip = img_clip.set_audio(final_audio)
                output_filename = (f"{base_filename}.mp4" if total_chunks == 1
                                   else f"{base_filename}_Part_{current_chunk_index}.mp4")
                output_path = os.path.join(self.output_dir, output_filename)
                img_clip.write_videofile(output_path, fps=1, codec="libx264", audio_codec="aac", logger=None)
                for clip in audio_clips: clip.close()
                final_audio.close()
                img_clip.close()
            except Exception as e:
                print(f"Error processing chunk {current_chunk_index}: {e}")

        self.lbl_status.configure(text=L["status_done"])
        self.btn_start.configure(state="normal")
        self.entry_chunk.configure(state="normal")
        self.entry_filename.configure(state="normal")
        self.slider_font_size.configure(state="normal")
        messagebox.showinfo("Success", L["status_done"])


# ══════════════════════════════════════════════════════════════════════════════
#
#  ███╗   ██╗ ██████╗ ██╗   ██╗███████╗██╗         ██╗   ██╗██████╗
#  ████╗  ██║██╔═══██╗██║   ██║██╔════╝██║         ██║   ██║╚════██╗
#  ██╔██╗ ██║██║   ██║██║   ██║█████╗  ██║         ██║   ██║ █████╔╝
#  ██║╚██╗██║██║   ██║╚██╗ ██╔╝██╔══╝  ██║         ╚██╗ ██╔╝ ╚═══██╗
#  ██║ ╚████║╚██████╔╝ ╚████╔╝ ███████╗███████╗     ╚████╔╝ ██████╔╝
#  ╚═╝  ╚═══╝ ╚═════╝   ╚═══╝  ╚══════╝╚══════╝      ╚═══╝  ╚═════╝
#
#  novel_by_keawgood_V2.py — Auto Novel Downloader (Hybrid Mode)
# ══════════════════════════════════════════════════════════════════════════════

try:
    from curl_cffi import requests as cffi_requests
    HAS_CURL = True
except ImportError:
    import requests as cffi_requests  # type: ignore
    HAS_CURL = False

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

NV_LANG: dict = {
    "th": {
        "app_title": "📖 Novel By Keawgood", "subtitle": "โหลดนิยายอัตโนมัติ (Hybrid Mode)",
        "url_label": "URL หน้าแนะนำนิยาย หรือ สารบัญ", "url_placeholder": "ตัวอย่าง: https://twkan.com/book/61470.html",
        "fetch_btn": "🔍 ดึงรายการตอน", "fetching": "⏳ กำลังดึงข้อมูล…",
        "chapters_found": "พบทั้งหมด: {n} ตอน", "no_chapters": "— ยังไม่พบตอน —",
        "range_label": "ช่วงตอนที่ต้องการโหลด", "from_label": "จากตอนที่:", "to_label": "ถึงตอนที่:",
        "workers_label": "Workers (เธรด):", "delay_min": "Delay ต่ำสุด (วิ):", "delay_max": "Delay สูงสุด (วิ):",
        "headless_label": "ซ่อนหน้าต่าง Chrome", "naming_label": "รูปแบบชื่อไฟล์",
        "naming_hint": "ใช้ [n] แทนเลขตอน — เช่น  ตอนที่ [n]  หรือ  Chapter [n]",
        "save_label": "โฟลเดอร์บันทึก:", "browse_btn": "📁 เลือก",
        "start_btn": "⚡ เริ่มโหลด", "stop_btn": "⛔ หยุด", "clear_btn": "🗑 ล้าง Log",
        "log_label": "บันทึกกิจกรรม", "bypass_btn": "🛡 ขอ Cloudflare Cookie",
        "tip_workers": "💡 หมายเหตุ: ระบบป้องกันการข้ามตอน หากโหลดไม่ผ่านจะพยายามโหลดตอนเดิมซ้ำจนสำเร็จ",
        "lang_btn": "🌐 English", "theme_btn": "☀️ โหมดสว่าง",
        "err_no_url": "กรุณากรอก URL ก่อน", "err_no_fetch": "กรุณากด 'ดึงรายการตอน' ก่อนเริ่มโหลด",
        "err_no_dir": "กรุณาเลือกโฟลเดอร์บันทึก",
        "err_range": "ช่วงตอนไม่ถูกต้อง (ต้องอยู่ระหว่าง 1–{n})", "err_workers": "Workers ต้องเป็น 1–20",
        "err_delay": "Delay ต้องเป็นตัวเลขบวก และ min ≤ max",
        "done_msg": "✨ โหลดสำเร็จครบถ้วน! สำเร็จ {ok} | {t:.1f}s",
        "stopped_msg": "⛔ หยุดโดยผู้ใช้ — สำเร็จ {ok} | {t:.1f}s",
        "saved_at": "📂 ไฟล์บันทึกที่: {d}", "paste": "📋 วาง", "copy": "📄 คัดลอก",
    },
    "en": {
        "app_title": "📖 Novel By Keawgood", "subtitle": "Automatic Novel Downloader (Hybrid Mode)",
        "url_label": "Novel Info URL or TOC", "url_placeholder": "Example: https://twkan.com/book/61470.html",
        "fetch_btn": "🔍 Fetch Chapters", "fetching": "⏳ Fetching…",
        "chapters_found": "Found: {n} chapters", "no_chapters": "— No chapters found —",
        "range_label": "Chapter Range to Download", "from_label": "From chapter:", "to_label": "To chapter:",
        "workers_label": "Workers (threads):", "delay_min": "Delay min (s):", "delay_max": "Delay max (s):",
        "headless_label": "Hide Chrome window", "naming_label": "File Naming Pattern",
        "naming_hint": "Use [n] for chapter number — e.g.  Chapter [n]  or  ตอนที่ [n]",
        "save_label": "Save folder:", "browse_btn": "📁 Browse",
        "start_btn": "⚡ Start Download", "stop_btn": "⛔ Stop", "clear_btn": "🗑 Clear Log",
        "log_label": "Activity Log", "bypass_btn": "🛡 Get Cloudflare Cookie",
        "tip_workers": "💡 Note: Failed chapters will be retried infinitely until successful.",
        "lang_btn": "🌐 ภาษาไทย", "theme_btn": "☀️ Light Mode",
        "err_no_url": "Please enter a URL first", "err_no_fetch": "Please fetch chapters before starting download",
        "err_no_dir": "Please select a save folder",
        "err_range": "Invalid chapter range (must be 1–{n})", "err_workers": "Workers must be between 1–20",
        "err_delay": "Delay must be positive numbers and min ≤ max",
        "done_msg": "✨ Done! Success {ok} | {t:.1f}s", "stopped_msg": "⛔ Stopped by user — Success {ok} | {t:.1f}s",
        "saved_at": "📂 Files saved at: {d}", "paste": "📋 Paste", "copy": "📄 Copy",
    },
}

NV_THEMES: dict = {
    "dark": {
        "BG": "#0f111a", "PANEL": "#1a1d2d", "ACCENT": "#6a5acd",
        "ACCENT_H": "#7b68ee", "ACCENT2": "#20c997", "FG": "#e2e8f0",
        "FG_DIM": "#94a3b8", "SUCCESS": "#20c997", "ERROR": "#f43f5e",
        "WARN": "#f59e0b", "ENTRY_BG": "#0b0c13", "BORDER": "#2e344e",
    },
    "light": {
        "BG": "#f8fafc", "PANEL": "#ffffff", "ACCENT": "#4f46e5",
        "ACCENT_H": "#6366f1", "ACCENT2": "#059669", "FG": "#0f172a",
        "FG_DIM": "#64748b", "SUCCESS": "#059669", "ERROR": "#e11d48",
        "WARN": "#d97706", "ENTRY_BG": "#f1f5f9", "BORDER": "#e2e8f0",
    },
}

_nv_cf_lock = threading.Lock()
_nv_cf_cookies: list = []
_nv_cf_user_agent: str = ""
NV_USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
]

NV_SITE_RULES: dict = {
    "twkan.com": {
        "link_filter": lambda abs_url, bid: re.search(r'\d+\.html$', abs_url) and "index.html" not in abs_url,
        "content_selector": ("div", {"id": re.compile(r'(content|chaptercontent|txtcontent0|BookText|read)', re.I)}),
        "unwanted_tags": ["script", "style", "ins", "h1"],
        "encoding": "utf-8",
    },
    "default": {
        "link_filter": lambda abs_url, bid: len(abs_url) > 10 and not abs_url.endswith("#") and (
            re.search(r'(chapter|chap|vol|part|/p/|/read/|/txt/|\d{3,}\.html|\d+/?$)', abs_url.lower())
            or (bid and bid in abs_url)),
        "content_selector": ("div", {"class": re.compile(r'(content|chapter|text|read|body|main|entry|post|article)', re.I)}),
        "unwanted_tags": ["script", "style", "ins", "aside", "nav", "header", "footer", "iframe"],
        "encoding": None,
    },
}

_NV_STEALTH_JS = """
Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
Object.defineProperty(navigator, 'languages', { get: () => ['en-US', 'en'] });
Object.defineProperty(navigator, 'platform', { get: () => 'Win32' });
"""

_NV_AD_PATTERNS = re.compile(
    r'(advertisement|sponsored|subscribe|follow us|please support|read at|'
    r'translator.?note|t/n:|tl.?note|visit .{0,30} for .{0,30} chapters|'
    r'patreon\.com|discord\.gg|ko-?fi\.com|www\.\w+\.com|'
    r'手机版阅读|最新网址|更新最快|下载APP|无弹窗|顶点小说|QQ群|微信号|公众号|一秒记住|'
    r'推荐本书|txt下载|求推荐|求收藏|章节错误|点此举报|加入书签|上一章|返回目录|下一章|'
    r'天才一秒记住本站地址|twkan|台湾看小说|台灣小說網|₮₩₭₳₦|請記住|觀看最快|章節更新|loadAdv)',
    re.I
)


def nv_playwright_get_cookies(url: str, log_fn, headless: bool = True) -> bool:
    global _nv_cf_cookies, _nv_cf_user_agent
    if not HAS_PLAYWRIGHT:
        log_fn("⚠️ Playwright not installed — skipping browser bypass"); return False
    log_fn("🌐 Opening stealth browser to obtain Cloudflare clearance…")
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=headless, args=["--disable-blink-features=AutomationControlled"])
            ctx = browser.new_context(user_agent=random.choice(NV_USER_AGENTS), viewport={"width": 1366, "height": 768})
            ctx.add_init_script(_NV_STEALTH_JS)
            page = ctx.new_page()
            try:
                page.goto(url, wait_until="domcontentloaded", timeout=30_000)
            except Exception:
                pass
            for _ in range(20):
                if "just a moment" not in page.title().lower() and "cloudflare" not in page.title().lower():
                    break
                time.sleep(1)
            _nv_cf_cookies = ctx.cookies()
            _nv_cf_user_agent = page.evaluate("navigator.userAgent")
            browser.close()
        log_fn("✅ Cookie obtained!" if _nv_cf_cookies else "❌ No cookies obtained")
        return bool(_nv_cf_cookies)
    except Exception as e:
        log_fn(f"❌ Playwright error: {e}"); return False


def nv_make_session():
    ua = _nv_cf_user_agent or random.choice(NV_USER_AGENTS)
    sess = cffi_requests.Session(impersonate="chrome120") if HAS_CURL else cffi_requests.Session()
    sess.headers.update({"User-Agent": ua, "Accept-Language": "en-US,en;q=0.9"})
    for c in _nv_cf_cookies:
        try:
            sess.cookies.set(c["name"], c["value"], domain=c.get("domain", ""))
        except Exception:
            pass
    return sess


def nv_decode_response(content: bytes, hint=None) -> str:
    if hint:
        try: return content.decode(hint, errors="replace")
        except Exception: pass
    if HAS_CHARDET:
        enc = chardet.detect(content[:4096]).get("encoding") or "utf-8"
        try: return content.decode(enc, errors="replace")
        except Exception: pass
    for enc in ("utf-8", "gbk", "big5"):
        try: return content.decode(enc, errors="replace")
        except Exception: continue
    return content.decode("utf-8", errors="replace")


def nv_detect_site(url: str) -> dict:
    for domain, rules in NV_SITE_RULES.items():
        if domain != "default" and domain in url:
            return rules
    return NV_SITE_RULES["default"]


def nv_extract_book_id(url: str) -> str:
    for pat in [r'/book/(\w+)', r'/b/(\w+)', r'/n/(\w+)', r'/(\d{4,})/?(?:\?|$|#)']:
        m = re.search(pat, url)
        if m: return m.group(1)
    parts = url.rstrip("/").split("/")
    return re.sub(r'[^\w]', '', parts[-1] if parts else "") or ""


def nv_fetch_toc(catalog_url: str, log_fn=print) -> list:
    if not HAS_BS4:
        log_fn("❌ beautifulsoup4 ไม่ได้ติดตั้ง — pip install beautifulsoup4"); return []
    rules = nv_detect_site(catalog_url)
    book_id = nv_extract_book_id(catalog_url)
    session = nv_make_session()
    html_text = None
    for _ in range(3):
        try:
            resp = session.get(catalog_url, timeout=20)
            if resp.status_code == 200:
                html_text = nv_decode_response(resp.content, rules.get("encoding")); break
        except Exception:
            time.sleep(2)
    if html_text is None:
        with _nv_cf_lock:
            if not _nv_cf_cookies:
                nv_playwright_get_cookies(catalog_url, log_fn)
        session = nv_make_session()
        for attempt in range(4):
            try:
                resp = session.get(catalog_url, timeout=20)
                if resp.status_code == 200:
                    html_text = nv_decode_response(resp.content, rules.get("encoding")); break
                time.sleep(2 ** attempt)
            except Exception:
                time.sleep(2)
        else:
            return []
    if not html_text: return []
    soup = BeautifulSoup(html_text, "html.parser")
    links = []
    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        abs_url = urljoin(catalog_url, href)
        try:
            if rules["link_filter"](abs_url, book_id):
                links.append(abs_url)
        except Exception:
            pass
    unique = list(dict.fromkeys(links))
    try:
        unique.sort(key=lambda u: int(re.search(r'/(\d+)(?:\.html?)?/?(?:\?.*)?$', u).group(1)))
    except Exception:
        pass
    if not unique:
        log_fn("⚠️ หาลิงก์สารบัญจากหน้านี้ไม่พบ (อาจต้องใช้โหมด 'ไต่ลิงก์ทีละตอน')")
        return [catalog_url]
    return unique


def nv_clean_text(div, unwanted_tags: list) -> str:
    for bad in div.find_all(unwanted_tags):
        bad.decompose()
    html_content = re.sub(r'<br\s*/?>', '\n', str(div))
    clean_soup = BeautifulSoup(html_content, "html.parser")
    lines = []
    for ln in clean_soup.get_text(separator="\n").splitlines():
        s = ln.strip()
        if s and not _NV_AD_PATTERNS.search(s):
            lines.append(s)
    return "\n".join(lines).strip()


def nv_fetch_chapter(session, chapter_num, url, rules, save_dir, stop_event, naming_pattern, log_fn, dmin, dmax) -> str:
    attempt = 0
    while not stop_event.is_set():
        if attempt > 0:
            stop_event.wait(min(30, (2 ** min(attempt, 5))) + random.uniform(1.0, 3.0))
        else:
            stop_event.wait(random.uniform(dmin, dmax))
        if stop_event.is_set(): return f"⛔ [{chapter_num:04d}] ยกเลิก"
        attempt += 1
        try:
            resp = session.get(url, timeout=30, headers={"Referer": url})
            if resp.status_code == 200:
                soup = BeautifulSoup(nv_decode_response(resp.content, rules.get("encoding")), "html.parser")
                title = soup.find("h1")
                title_text = title.get_text(strip=True) if title else f"Chapter {chapter_num:04d}"
                tag, attrs = rules["content_selector"]
                div = soup.find(tag, attrs)
                if not div:
                    for pid in ["content", "txtcontent0", "chaptercontent", "BookText", "chapter-content", "txtContent"]:
                        div = soup.find("div", id=re.compile(pid, re.I))
                        if div: break
                if not div:
                    all_divs = soup.find_all("div")
                    if all_divs:
                        best_div = max(all_divs, key=lambda d: len(d.get_text(strip=True)))
                        if len(best_div.get_text(strip=True)) > 200: div = best_div
                if not div:
                    log_fn(f"⚠️ [{chapter_num:04d}] ไม่พบเนื้อหา จะพยายามโหลดใหม่..."); continue
                text = nv_clean_text(div, rules.get("unwanted_tags", []))
                if len(text) < 30:
                    log_fn(f"⚠️ [{chapter_num:04d}] เนื้อหาสั้นผิดปกติ จะพยายามโหลดใหม่..."); continue
                safe_title = re.sub(r'[\\/*?:"<>|\r\n]', "", title_text)[:80]
                name = (re.sub(r'[\\/*?:"<>|]', "", naming_pattern.replace("[n]", str(chapter_num)))
                        if naming_pattern.strip() else f"{chapter_num:04d}_{safe_title}")
                with open(os.path.join(save_dir, name.strip() + ".txt"), "w", encoding="utf-8") as f:
                    f.write(f"{title_text}\n\n{text}\n\n(本集结束)")
                return f"✅ [{chapter_num:04d}] {title_text[:60]}"
            elif resp.status_code in (403, 429, 503):
                log_fn(f"⚠️ [{chapter_num:04d}] ติด Block ({resp.status_code}) จะพยายามโหลดใหม่...")
            else:
                log_fn(f"⚠️ [{chapter_num:04d}] HTTP {resp.status_code} จะพยายามโหลดใหม่...")
        except Exception:
            log_fn(f"⚠️ [{chapter_num:04d}] Error: เชื่อมต่อขัดข้อง จะพยายามโหลดใหม่...")
    return f"⛔ [{chapter_num:04d}] ยกเลิก"


class NovelByKeawgoodWindow(tk.Toplevel):
    """Novel V2 — runs as a Toplevel inside the CTk hub."""

    def __init__(self, master):
        super().__init__(master)
        self._lang       = "th"
        self._theme_name = "dark"
        self._theme      = NV_THEMES["dark"]
        self._links: list = []
        self._running    = False
        self._stop_event = threading.Event()

        self.title(NV_LANG["th"]["app_title"])
        self.geometry("950x880")
        self.minsize(800, 680)
        self.resizable(True, True)

        self.font_family = "TH Sarabun PSK"
        self.font_h1   = tkfont.Font(family=self.font_family, size=28, weight="bold")
        self.font_h2   = tkfont.Font(family=self.font_family, size=20, weight="bold")
        self.font_body = tkfont.Font(family=self.font_family, size=18)
        self.font_btn  = tkfont.Font(family=self.font_family, size=18, weight="bold")
        self.font_log  = tkfont.Font(family=self.font_family, size=16)

        self._apply_theme()
        self._setup_global_clipboard_bindings()
        self._build()
        self.lift()
        self.focus_force()

    def _adjust_font_size(self, delta):
        for f in [self.font_h1, self.font_h2, self.font_body, self.font_btn, self.font_log]:
            f.configure(size=max(10, f.cget("size") + delta))

    def _setup_global_clipboard_bindings(self):
        self.bind_all("<KeyPress>", self._handle_global_shortcuts)

    def _handle_global_shortcuts(self, event):
        if not (event.state & 0x0004 or event.state & 0x0008 or event.state & 0x20000): return
        char = getattr(event, "char", "").lower()
        keysym = getattr(event, "keysym", "").lower()
        if char in ("c", "แ") or keysym in ("c", "oae"): return self._execute_copy(event)
        elif char in ("v", "อ") or keysym in ("v", "oang"): return self._execute_paste(event)
        elif char in ("x", "ป") or keysym in ("x", "porpla"): return self._execute_cut(event)
        elif char in ("a", "ฟ") or keysym in ("a", "forfan"): return self._execute_select_all(event)

    def _execute_copy(self, event):
        w = event.widget
        try:
            text = w.selection_get() if isinstance(w, (tk.Entry, ttk.Entry)) else w.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.clipboard_clear(); self.clipboard_append(text); return "break"
        except tk.TclError: pass

    def _execute_paste(self, event):
        w = event.widget
        try:
            text = self.clipboard_get()
            if isinstance(w, (tk.Entry, ttk.Entry)):
                try: w.delete(tk.SEL_FIRST, tk.SEL_LAST)
                except tk.TclError: pass
                w.insert(tk.INSERT, text); return "break"
            elif isinstance(w, tk.Text) and w.cget("state") == "normal":
                w.insert(tk.INSERT, text); return "break"
        except tk.TclError: pass

    def _execute_cut(self, event):
        w = event.widget
        try:
            if isinstance(w, (tk.Entry, ttk.Entry)):
                text = w.selection_get()
                self.clipboard_clear(); self.clipboard_append(text)
                w.delete(tk.SEL_FIRST, tk.SEL_LAST); return "break"
        except tk.TclError: pass

    def _execute_select_all(self, event):
        w = event.widget
        if isinstance(w, (tk.Entry, ttk.Entry)):
            w.select_range(0, tk.END); w.icursor(tk.END); return "break"
        elif isinstance(w, tk.Text):
            w.tag_add(tk.SEL, "1.0", tk.END); return "break"

    def t(self, key: str, **kw) -> str:
        s = NV_LANG[self._lang].get(key, key)
        return s.format(**kw) if kw else s

    @property
    def T(self): return self._theme

    def _apply_theme(self):
        self._theme = NV_THEMES[self._theme_name]
        self.configure(bg=self.T["BG"])

    def _toggle_theme(self):
        self._theme_name = "light" if self._theme_name == "dark" else "dark"
        self._apply_theme()
        for w in self.winfo_children(): w.destroy()
        self._build()

    def _toggle_lang(self):
        self._lang = "en" if self._lang == "th" else "th"
        for w in self.winfo_children(): w.destroy()
        self._build()

    def _ttk_style(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        s.configure("Accent.Horizontal.TProgressbar", troughcolor=self.T["ENTRY_BG"],
                    background=self.T["ACCENT2"], bordercolor=self.T["ENTRY_BG"], relief="flat")

    def _card(self, parent):
        return tk.Frame(parent, bg=self.T["PANEL"], bd=0, highlightthickness=1,
                        highlightbackground=self.T["BORDER"], padx=20, pady=16)

    def _btn(self, parent, key, cmd, bg=None, fg=None, hover_bg=None, width=None, text=None, custom_font=None):
        label = text if text is not None else self.t(key)
        btn = tk.Button(parent, text=label, command=cmd, bg=bg or self.T["ACCENT"], fg=fg or "white",
                        font=custom_font or self.font_btn, relief="flat", cursor="hand2",
                        activebackground=hover_bg or self.T["ACCENT_H"], activeforeground=fg or "white",
                        padx=12, pady=6, bd=0)
        if width: btn.config(width=width)
        return btn

    def _entry(self, parent, textvariable, width=None, justify="left", ipady=4):
        kw = dict(textvariable=textvariable, bg=self.T["ENTRY_BG"], fg=self.T["FG"],
                  insertbackground=self.T["FG"], relief="flat", font=self.font_body, bd=6, justify=justify)
        if width: kw["width"] = width
        return tk.Entry(parent, **kw)

    def _label(self, parent, key, font=None, fg=None, text=None):
        return tk.Label(parent, text=text if text is not None else self.t(key),
                        font=font or self.font_body, bg=self.T["PANEL"], fg=fg or self.T["FG"])

    def _build(self):
        self._ttk_style()
        hdr = tk.Frame(self, bg=self.T["BG"])
        hdr.pack(fill="x", padx=24, pady=(15, 5))
        left = tk.Frame(hdr, bg=self.T["BG"])
        left.pack(side="left", fill="y")
        tk.Label(left, text=self.t("app_title"), font=self.font_h1, bg=self.T["BG"], fg=self.T["ACCENT"]).pack(anchor="w")
        tk.Label(left, text=self.t("subtitle"), font=self.font_body, bg=self.T["BG"], fg=self.T["FG_DIM"]).pack(anchor="w")
        right = tk.Frame(hdr, bg=self.T["BG"])
        right.pack(side="right", fill="y")
        ctrl_row = tk.Frame(right, bg=self.T["BG"])
        ctrl_row.pack(anchor="e")
        tk.Button(ctrl_row, text="A-", command=lambda: self._adjust_font_size(-2), bg=self.T["BORDER"], fg=self.T["FG"],
                  font=("Arial", 10, "bold"), relief="flat", cursor="hand2", padx=8, pady=4, bd=0).pack(side="left", padx=(0, 4))
        tk.Button(ctrl_row, text="A+", command=lambda: self._adjust_font_size(2), bg=self.T["BORDER"], fg=self.T["FG"],
                  font=("Arial", 10, "bold"), relief="flat", cursor="hand2", padx=8, pady=4, bd=0).pack(side="left", padx=(0, 12))
        theme_icon = "🌙" if self._theme_name == "dark" else "☀️"
        tk.Button(ctrl_row, text=theme_icon, command=self._toggle_theme, bg=self.T["BORDER"], fg=self.T["FG"],
                  font=("Arial", 10), relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left", padx=(0, 8))
        tk.Button(ctrl_row, text=self.t("lang_btn"), command=self._toggle_lang, bg=self.T["BORDER"], fg=self.T["FG"],
                  font=self.font_btn, relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left")

        uc = self._card(self)
        uc.pack(fill="x", padx=24, pady=8)
        uc.columnconfigure(0, weight=1)
        self._label(uc, "url_label", font=self.font_h2).grid(row=0, column=0, sticky="w", pady=(0, 4))
        ur = tk.Frame(uc, bg=self.T["PANEL"])
        ur.grid(row=1, column=0, sticky="ew")
        ur.columnconfigure(0, weight=1)
        self.url_var = tk.StringVar(value="https://twkan.com/book/61470.html")
        self.url_entry = self._entry(ur, self.url_var, ipady=6)
        self.url_entry.grid(row=0, column=0, sticky="ew")
        btn_action = tk.Frame(ur, bg=self.T["PANEL"])
        btn_action.grid(row=0, column=1, padx=(10, 10))
        tk.Button(btn_action, text=self.t("paste"), command=self._ui_paste_url,
                  bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_body,
                  relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left", padx=2)
        tk.Button(btn_action, text=self.t("copy"), command=self._ui_copy_url,
                  bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_body,
                  relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left", padx=2)
        self.fetch_btn = self._btn(ur, "fetch_btn", self._on_fetch)
        self.fetch_btn.grid(row=0, column=2)
        self.count_lbl = tk.Label(uc, text=self.t("no_chapters"), font=self.font_btn,
                                  bg=self.T["PANEL"], fg=self.T["ACCENT2"])
        self.count_lbl.grid(row=2, column=0, sticky="w", pady=(8, 0))
        mr = tk.Frame(uc, bg=self.T["PANEL"])
        mr.grid(row=3, column=0, sticky="w", pady=(5, 0))
        self.scrape_mode = tk.StringVar(value="crawler")
        tk.Label(mr, text="ระบบดึงข้อมูล:", font=self.font_btn, bg=self.T["PANEL"], fg=self.T["FG"]).pack(side="left")
        tk.Radiobutton(mr, text="ดึงจากสารบัญ (โหลดพร้อมกันหลายไฟล์)", variable=self.scrape_mode, value="concurrent",
                       font=self.font_body, bg=self.T["PANEL"], fg=self.T["ACCENT2"],
                       selectcolor=self.T["ENTRY_BG"], activebackground=self.T["PANEL"]).pack(side="left", padx=10)
        tk.Radiobutton(mr, text="ไต่ลิงก์ทีละตอน (🌟ชัวร์ 100% สำหรับ TWKAN)", variable=self.scrape_mode, value="crawler",
                       font=self.font_body, bg=self.T["PANEL"], fg=self.T["WARN"],
                       selectcolor=self.T["ENTRY_BG"], activebackground=self.T["PANEL"]).pack(side="left", padx=10)

        cc = self._card(self)
        cc.pack(fill="x", padx=24, pady=8)
        for c in range(8): cc.columnconfigure(c, weight=1)
        rows_cfg = [
            ("from_label",    "start_var",     "1",   0, 0),
            ("to_label",      "end_var",        "10",  0, 2),
            ("workers_label", "workers_var",   "5",   0, 4),
            ("delay_min",     "delay_min_var", "1.5", 1, 0),
            ("delay_max",     "delay_max_var", "4.0", 1, 2),
        ]
        for lkey, attr, default, row, col in rows_cfg:
            self._label(cc, lkey).grid(row=row, column=col, sticky="w", pady=(0 if row == 0 else 12, 0))
            var = tk.StringVar(value=default)
            setattr(self, attr, var)
            self._entry(cc, var, width=6 if attr == "workers_var" else 8, justify="center"
                        ).grid(row=row, column=col+1, padx=(4, 16), ipady=2, pady=(0 if row == 0 else 12, 0))
        self.headless_var = tk.BooleanVar(value=True)
        tk.Checkbutton(cc, text=self.t("headless_label"), variable=self.headless_var,
                       font=self.font_body, bg=self.T["PANEL"], fg=self.T["FG_DIM"],
                       selectcolor=self.T["ENTRY_BG"], activebackground=self.T["PANEL"]
                       ).grid(row=0, column=6, columnspan=2, sticky="w", padx=(8, 0))
        self._btn(cc, "bypass_btn", self._on_bypass, bg=self.T["WARN"], fg="#1a1d2e", hover_bg="#fcd34d"
                  ).grid(row=1, column=4, columnspan=4, sticky="e", pady=(12, 0))
        tk.Frame(cc, bg=self.T["BORDER"], height=1).grid(row=2, column=0, columnspan=8, sticky="ew", pady=(16, 12))
        self._label(cc, "naming_label", font=self.font_h2).grid(row=3, column=0, columnspan=8, sticky="w")
        self.naming_var = tk.StringVar(value="ตอนที่ [n]" if self._lang == "th" else "Chapter [n]")
        nr = tk.Frame(cc, bg=self.T["PANEL"])
        nr.grid(row=4, column=0, columnspan=8, sticky="ew", pady=(8, 0))
        nr.columnconfigure(0, weight=1)
        self._entry(nr, self.naming_var, ipady=4).grid(row=0, column=0, sticky="ew")
        for i, preset in enumerate(["ตอนที่ [n]", "Chapter [n]", "第[n]章"]):
            tk.Button(nr, text=preset, command=lambda p=preset: self.naming_var.set(p),
                      bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_body,
                      relief="flat", cursor="hand2", padx=10, pady=4, bd=0).grid(row=0, column=i+1, padx=(8, 0))
        tk.Frame(cc, bg=self.T["BORDER"], height=1).grid(row=5, column=0, columnspan=8, sticky="ew", pady=(16, 12))
        dr = tk.Frame(cc, bg=self.T["PANEL"])
        dr.grid(row=6, column=0, columnspan=8, sticky="ew")
        dr.columnconfigure(1, weight=1)
        self._label(dr, "save_label").grid(row=0, column=0, sticky="w")
        self.save_dir_var = tk.StringVar(value=str(Path.home() / "Downloads" / "novels"))
        self._entry(dr, self.save_dir_var, ipady=4).grid(row=0, column=1, sticky="ew", padx=(10, 10))
        self._btn(dr, "browse_btn", self._browse, bg=self.T["ENTRY_BG"], fg=self.T["FG"]).grid(row=0, column=2)
        tk.Label(cc, text=self.t("tip_workers"), font=self.font_body, bg=self.T["PANEL"],
                 fg=self.T["WARN"]).grid(row=7, column=0, columnspan=8, sticky="w", pady=(10, 0))

        pg = tk.Frame(self, bg=self.T["BG"])
        pg.pack(fill="x", padx=24, pady=(8, 4))
        pg.columnconfigure(0, weight=1)
        self.prog_var = tk.DoubleVar(value=0)
        ttk.Progressbar(pg, variable=self.prog_var, maximum=100,
                        style="Accent.Horizontal.TProgressbar").grid(row=0, column=0, sticky="ew")
        self.prog_lbl = tk.Label(pg, text="0 / 0", font=self.font_btn, bg=self.T["BG"], fg=self.T["FG_DIM"])
        self.prog_lbl.grid(row=0, column=1, padx=(12, 0))

        br = tk.Frame(self, bg=self.T["BG"])
        br.pack(pady=8)
        self.start_btn = self._btn(br, "start_btn", self._on_start, bg=self.T["ACCENT2"], fg="#0f111a",
                                   hover_bg="#34d399", width=18)
        self.start_btn.pack(side="left", padx=8)
        self.stop_btn = self._btn(br, "stop_btn", self._on_stop, bg=self.T["ERROR"], fg="white",
                                  hover_bg="#fb7185", width=10)
        self.stop_btn.pack(side="left", padx=8)
        self.stop_btn.config(state="disabled")
        self._btn(br, "clear_btn", self._clear_log, bg=self.T["BORDER"], fg=self.T["FG"],
                  hover_bg=self.T["ENTRY_BG"], width=10).pack(side="left", padx=8)

        lf = self._card(self)
        lf.pack(fill="both", expand=True, padx=24, pady=(4, 20))
        lf.columnconfigure(0, weight=1)
        lf.rowconfigure(1, weight=1)
        self._label(lf, "log_label", font=self.font_h2).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.log = tk.Text(lf, bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_log,
                           relief="flat", wrap="word", state="disabled", bd=10)
        self.log.grid(row=1, column=0, sticky="nsew")
        sb = tk.Scrollbar(lf, command=self.log.yview, bg=self.T["PANEL"], troughcolor=self.T["ENTRY_BG"])
        sb.grid(row=1, column=1, sticky="ns", padx=(4, 0))
        self.log.config(yscrollcommand=sb.set)
        for tag, col in [("ok", self.T["SUCCESS"]), ("err", self.T["ERROR"]),
                         ("warn", self.T["WARN"]), ("info", self.T["FG_DIM"]), ("head", self.T["ACCENT_H"])]:
            self.log.tag_config(tag, foreground=col)
        self._log("Novel By Keawgood (Hybrid TWKAN Edition) — ready ✓", "head")

    def _ui_paste_url(self):
        try:
            text = self.clipboard_get()
            self.url_entry.delete(0, tk.END)
            self.url_entry.insert(0, text)
        except tk.TclError: pass

    def _ui_copy_url(self):
        text = self.url_entry.get()
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)

    def _log(self, msg: str, tag: str = "info"):
        def _do():
            self.log.config(state="normal")
            self.log.insert("end", f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n", tag)
            self.log.see("end")
            self.log.config(state="disabled")
        self.after(0, _do)

    def _clear_log(self):
        self.log.config(state="normal")
        self.log.delete("1.0", "end")
        self.log.config(state="disabled")

    def _browse(self):
        if d := filedialog.askdirectory(title="Select Folder"):
            self.save_dir_var.set(d)

    def _on_bypass(self):
        if not (url := self.url_var.get().strip()):
            return messagebox.showerror("Error", self.t("err_no_url"))
        self._log("🛡 Requesting Cloudflare clearance manually…", "warn")
        threading.Thread(target=lambda: nv_playwright_get_cookies(
            url, lambda m: self._log(m), self.headless_var.get()), daemon=True).start()

    def _on_fetch(self):
        if not (url := self.url_var.get().strip()):
            return messagebox.showerror("Error", self.t("err_no_url"))
        if self.scrape_mode.get() == "crawler":
            self._log("💡 อยู่ใน 'โหมดไต่ลิงก์' ไม่จำเป็นต้องดึงสารบัญ สามารถกดปุ่ม [เริ่มโหลด] ได้เลย", "head")
            return
        self.fetch_btn.config(state="disabled", text=self.t("fetching"))
        self._log(f"🔍 {self.t('fetch_btn')}: {url}", "head")

        def _worker():
            self._links = nv_fetch_toc(url, log_fn=lambda m: self._log(m, "warn"))
            def _upd():
                n = len(self._links)
                self.count_lbl.config(text=self.t("chapters_found", n=n) if n else self.t("no_chapters"),
                                      fg=self.T["SUCCESS"] if n else self.T["ERROR"])
                if n: self.end_var.set(str(n))
                self.fetch_btn.config(state="normal", text=self.t("fetch_btn"))
            self.after(0, _upd)
        threading.Thread(target=_worker, daemon=True).start()

    def _on_start(self):
        if self._running: return
        if not (save_dir := self.save_dir_var.get().strip()):
            return messagebox.showerror("Error", self.t("err_no_dir"))
        os.makedirs(save_dir, exist_ok=True)
        mode = self.scrape_mode.get()
        if mode == "crawler":
            if not self.url_var.get().strip():
                return messagebox.showerror("Error", self.t("err_no_url"))
            self._running = True
            self._stop_event.clear()
            self.start_btn.config(state="disabled")
            self.stop_btn.config(state="normal")
            threading.Thread(target=self._crawler_worker, daemon=True).start()
            return
        if not self._links:
            return messagebox.showerror("Error", self.t("err_no_fetch"))
        try:
            s = int(self.start_var.get())
            e = int(self.end_var.get())
            workers = int(self.workers_var.get())
            dmin = float(self.delay_min_var.get())
            dmax = float(self.delay_max_var.get())
            assert 1 <= s <= e <= len(self._links) and 1 <= workers <= 20 and 0 <= dmin <= dmax
        except Exception:
            return messagebox.showerror("Input Error", "Check numbers (range, workers, delays)")
        self._running = True
        naming = self.naming_var.get().strip()
        self._stop_event.clear()
        self.start_btn.config(state="disabled")
        self.stop_btn.config(state="normal")
        selected = self._links[s - 1: e]
        self.prog_var.set(0)
        self.prog_lbl.config(text=f"0 / {len(selected)}")
        rules = nv_detect_site(self.url_var.get())
        self._log(f"⚡ Downloading {len(selected)} chapters (Concurrent Mode)...", "head")

        def _worker():
            session = nv_make_session()
            completed = 0
            t0 = time.time()
            log_wrapper = lambda m: self._log(m, "warn")
            with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as ex:
                futs = {ex.submit(nv_fetch_chapter, session, s + i, u, rules, save_dir,
                                  self._stop_event, naming, log_wrapper, dmin, dmax): s + i
                        for i, u in enumerate(selected) if not self._stop_event.is_set()}
                for fut in concurrent.futures.as_completed(futs):
                    if self._stop_event.is_set(): break
                    res = fut.result()
                    if "✅" in res: completed += 1; self._log(res, "ok")
                    elif "⛔" in res: self._log(res, "warn")
                    self.after(0, lambda c=completed: (
                        self.prog_var.set(c / len(selected) * 100),
                        self.prog_lbl.config(text=f"{c} / {len(selected)}")))
            self.after(0, lambda: self._on_done(completed, time.time() - t0))
        threading.Thread(target=_worker, daemon=True).start()

    def _crawler_worker(self):
        url = self.url_var.get().strip()
        save_dir = self.save_dir_var.get().strip()
        naming = self.naming_var.get().strip()
        rules = nv_detect_site(url)
        session = nv_make_session()
        current_url = url
        try: chapter_num = int(self.start_var.get())
        except Exception: chapter_num = 1
        completed = 0
        try: dmin, dmax = float(self.delay_min_var.get()), float(self.delay_max_var.get())
        except Exception: dmin, dmax = 1.5, 4.0
        self._log("═" * 52, "head")
        self._log(f"⚡ เริ่ม [โหมดไต่ลิงก์] จาก: {url}", "head")
        self.after(0, lambda: (self.prog_lbl.config(text="โหมดไต่ลิงก์: กำลังเริ่มต้น..."), self.prog_var.set(0)))
        t0 = time.time()
        while not self._stop_event.is_set():
            attempt, success = 0, False
            while attempt < 5 and not self._stop_event.is_set():
                if attempt > 0:
                    self._stop_event.wait(min(30, (2 ** attempt)) + random.uniform(1, 3))
                else:
                    self._stop_event.wait(random.uniform(dmin, dmax))
                attempt += 1
                try:
                    if not HAS_BS4: self._log("❌ beautifulsoup4 ไม่ได้ติดตั้ง", "err"); self._stop_event.set(); break
                    resp = session.get(current_url, timeout=30, headers={"Referer": current_url})
                    if resp.status_code == 200:
                        soup = BeautifulSoup(nv_decode_response(resp.content, rules.get("encoding")), "html.parser")
                        tag, attrs = rules.get("content_selector", ("div", {}))
                        div = soup.find(tag, attrs)
                        if not div:
                            for pid in ["txtcontent0", "content", "chaptercontent", "BookText", "chapter-content", "txtContent"]:
                                div = soup.find("div", id=re.compile(pid, re.I))
                                if div: break
                        if not div:
                            all_divs = soup.find_all("div")
                            if all_divs:
                                best_div = max(all_divs, key=lambda d: len(d.get_text(strip=True)))
                                if len(best_div.get_text(strip=True)) > 200: div = best_div
                        if not div:
                            first_link = None
                            for a in soup.find_all("a", href=True):
                                href = a['href']
                                if re.search(r'\d+\.html$', href) and "index" not in href:
                                    first_link = urljoin(current_url, href); break
                            if first_link:
                                self._log(f"🧭 ตรวจพบว่าเป็นหน้าสารบัญ กระโดดไปตอนแรก: {first_link}", "info")
                                current_url = first_link; success = True; break
                            else:
                                self._log(f"⚠️ [{chapter_num:04d}] ไม่พบเนื้อหาและลิงก์ตอนแรก พยายามใหม่...", "warn"); continue
                        title = soup.find("h1")
                        title_text = title.get_text(strip=True) if title else f"Chapter {chapter_num:04d}"
                        text = nv_clean_text(div, rules.get("unwanted_tags", []))
                        if len(text) < 30:
                            self._log(f"⚠️ [{chapter_num:04d}] เนื้อหาสั้นผิดปกติ พยายามใหม่...", "warn"); continue
                        safe_title = re.sub(r'[\\/*?:"<>|\r\n]', "", title_text)[:80]
                        name = (re.sub(r'[\\/*?:"<>|]', "", naming.replace("[n]", str(chapter_num)))
                                if naming.strip() else f"{chapter_num:04d}_{safe_title}")
                        with open(os.path.join(save_dir, name.strip() + ".txt"), "w", encoding="utf-8") as f:
                            f.write(f"{title_text}\n\n{text}\n\n(本集结束)")
                        self._log(f"✅ [{chapter_num:04d}] {title_text[:60]}", "ok")
                        completed += 1; chapter_num += 1
                        self.after(0, lambda c=completed: self.prog_lbl.config(text=f"โหลดสะสม: {c} ตอน"))
                        next_a = soup.find("a", string=re.compile(r'下一章|下一页|Next|หน้าถัดไป'))
                        if not next_a:
                            self._log("🏁 ไม่พบปุ่ม 'หน้าถัดไป' สิ้นสุดการทำงาน", "ok"); self._stop_event.set(); break
                        next_url = urljoin(current_url, next_a['href'])
                        is_index = any(x in next_url.lower() for x in ['mulu', 'index', 'book', 'catalog', 'info'])
                        if is_index or next_url == current_url or next_url.endswith('/'):
                            self._log("🏁 ปุ่มหน้าถัดไปชี้กลับไปหน้าสารบัญ (โหลดถึงตอนล่าสุดแล้ว)", "ok")
                            self._stop_event.set(); break
                        current_url = next_url; success = True; break
                    elif resp.status_code in (403, 429, 503):
                        self._log(f"⚠️ [{chapter_num:04d}] ติด Block ({resp.status_code}) พยายามใหม่...", "warn")
                        if resp.status_code == 403 and attempt >= 2:
                            with _nv_cf_lock:
                                nv_playwright_get_cookies(current_url, lambda m: self._log(m, "warn"), self.headless_var.get())
                            session = nv_make_session()
                    else:
                        self._log(f"⚠️ [{chapter_num:04d}] HTTP {resp.status_code} พยายามใหม่...", "warn")
                except Exception:
                    self._log(f"⚠️ [{chapter_num:04d}] Error: การเชื่อมต่อขัดข้อง พยายามใหม่...", "warn")
            if not success and not self._stop_event.is_set():
                self._log(f"❌ โหลดตอนที่ {chapter_num} ซ้ำ 5 ครั้งไม่สำเร็จ โปรแกรมหยุดอัตโนมัติ", "err")
                self._stop_event.set(); break
        self.after(0, lambda: self._on_done(completed, time.time() - t0))

    def _on_done(self, completed, elapsed):
        self._running = False
        self.start_btn.config(state="normal")
        self.stop_btn.config(state="disabled")
        self._log(f"✨ ทำงานเสร็จสิ้น โหลดสำเร็จทั้งหมด {completed} ตอน ในเวลา {elapsed:.1f} วินาที", "ok")

    def _on_stop(self):
        self._stop_event.set()
        self._log("⛔ Stopping...", "warn")
        self.stop_btn.config(state="disabled")


# ══════════════════════════════════════════════════════════════════════════════
#
#  ███╗   ███╗ █████╗ ██╗███╗   ██╗    ██╗  ██╗██╗   ██╗██████╗
#  ████╗ ████║██╔══██╗██║████╗  ██║    ██║  ██║██║   ██║██╔══██╗
#  ██╔████╔██║███████║██║██╔██╗ ██║    ███████║██║   ██║██████╔╝
#  ██║╚██╔╝██║██╔══██║██║██║╚██╗██║    ██╔══██║██║   ██║██╔══██╗
#  ██║ ╚═╝ ██║██║  ██║██║██║ ╚████║    ██║  ██║╚██████╔╝██████╔╝
#  ╚═╝     ╚═╝╚═╝  ╚═╝╚═╝╚═╝  ╚═══╝   ╚═╝  ╚═╝ ╚═════╝ ╚═════╝
#
#  main_hub.py — Keawgood Universe Launcher
# ══════════════════════════════════════════════════════════════════════════════

class MainLauncher(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Keawgood Universe - หน้าต่างรวมโปรแกรม")
        self.geometry("450x580")
        self.resizable(False, False)
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        self._open_windows = {}
        self._build()

    def _build(self):
        ctk.CTkLabel(self, text="🌟 Keawgood Tools Hub",
                     font=("TH Sarabun PSK", 30, "bold")).pack(pady=(30, 5))
        ctk.CTkLabel(self, text="All-in-One Edition",
                     font=("TH Sarabun PSK", 16), text_color="gray").pack(pady=(0, 20))

        apps = [
            {"name": "📖 โหลดนิยายอัตโนมัติ (V2)", "key": "novel",  "color": "#e74c3c",
             "fn": lambda: NovelByKeawgoodWindow(self)},
            {"name": "🗂️ จัดการไฟล์นิยาย",          "key": "files",  "color": "#f39c12",
             "fn": lambda: ByKeawgoodWindow(self)},
            {"name": "✂️ จัดการคำศัพท์นิยาย",        "key": "vocab",  "color": "#27ae60",
             "fn": lambda: VocabOptimizerWindow(self)},
            {"name": "🎵 แปลงไฟล์เสียงเป็นวิดีโอ",   "key": "audio",  "color": "#8e44ad",
             "fn": lambda: AudioByKeawgoodWindow(self)},
        ]

        for app in apps:
            btn = ctk.CTkButton(
                self,
                text=app["name"],
                font=("TH Sarabun PSK", 20, "bold"),
                fg_color=app["color"],
                height=55,
                command=lambda a=app: self._open_app(a["key"], a["fn"])
            )
            btn.pack(pady=10, padx=40, fill="x")

        ctk.CTkLabel(self, text="คลิกเพื่อเปิดโปรแกรม | ทุกโปรแกรมเปิดได้พร้อมกัน",
                     text_color="gray", font=("TH Sarabun PSK", 14)).pack(side="bottom", pady=20)

    def _open_app(self, key: str, factory):
        """เปิดหน้าต่างใหม่ ถ้าเปิดอยู่แล้วให้ยก focus ขึ้นมา"""
        win = self._open_windows.get(key)
        if win and win.winfo_exists():
            win.lift()
            win.focus_force()
            return
        win = factory()
        self._open_windows[key] = win


if __name__ == "__main__":
    app = MainLauncher()
    app.mainloop()